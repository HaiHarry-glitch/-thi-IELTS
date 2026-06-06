import html
import json
import re
import sys
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = Path(r"D:\READ VOL 9-20260606T164451Z-3-001\READ VOL 9")
KEY_DIR = SOURCE_DIR / "KEY-EXPLANATION"
NORMALIZED_DIR = ROOT / "data" / "normalized"
NORMALIZED_INDEX = NORMALIZED_DIR / "_index.json"
FULL_TEST_INDEX = ROOT / "data" / "full-tests" / "_index.json"

BASE_ID = 909000
VOL = 9
ANSWER_OVERRIDES = {
    (3, 40): "YES",
}


def norm_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def html_p(text: str) -> str:
    return f"<p>{html.escape(text)}</p>"


def text_to_html(text: str) -> str:
    lines = [norm_text(line) for line in (text or "").splitlines()]
    return "".join(html_p(line) for line in lines if line)


def iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def iter_doc_units(doc: Document):
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = norm_text(block.text)
            if text:
                yield ("p", text)
        else:
            for row in block.rows:
                cells = [norm_text(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    yield ("cell", row_text)


def split_passages(doc_path: Path):
    doc = Document(doc_path)
    units = list(iter_doc_units(doc))
    starts = [i for i, (_kind, text) in enumerate(units) if re.match(r"^PASSAGE\s+[123]\b", text, re.I)]
    starts.append(len(units))
    passages = []
    for idx in range(len(starts) - 1):
        chunk = [text for _kind, text in units[starts[idx] + 1 : starts[idx + 1]]]
        passages.append(chunk)
    return passages


def get_title(chunk, passage_no: int) -> str:
    for i, text in enumerate(chunk):
        if re.match(r"Read the text and answer questions", text, re.I):
            if i + 1 < len(chunk):
                return chunk[i + 1]
    return f"Vol 9 Test Passage {passage_no}"


def passage_body(chunk):
    body = []
    for text in chunk:
        if re.match(r"Questions?\s+\d+", text, re.I):
            break
        if re.match(r"Read the text and answer questions", text, re.I):
            continue
        if re.fullmatch(r"[-–—_ ]{8,}", text):
            continue
        body.append(text)
    return body


def extract_question_texts(chunk):
    question_texts = {}
    in_questions = False
    for text in chunk:
        if re.match(r"Questions?\s+\d+", text, re.I):
            in_questions = True
            continue
        if not in_questions:
            continue
        if re.match(r"Read the text and answer questions", text, re.I):
            continue
        if re.match(r"List of Headings", text, re.I):
            continue
        if re.match(r"^[ivxlcdm]+\s+", text, re.I):
            continue
        if re.match(r"^[A-E]\s+", text):
            continue

        found = re.findall(r"(?<!\d)(\d{1,2})(?:\s*[-–]\s*(\d{1,2}))?", text)
        if not found:
            continue
        for start, end in found:
            a = int(start)
            b = int(end) if end else a
            if a < 1 or b > 40 or a > b:
                continue
            for order in range(a, b + 1):
                question_texts.setdefault(order, text)
    return question_texts


def extract_key_cells(key_path: Path):
    doc = Document(key_path)
    best = None
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 3:
            continue
        headers = [norm_text(cell.text).upper() for cell in table.rows[0].cells[:3]]
        if headers[:3] == ["PASSAGE 1", "PASSAGE 2", "PASSAGE 3"]:
            best = [table.rows[1].cells[i].text for i in range(3)]
    if not best:
        raise RuntimeError(f"No answer table found in {key_path}")
    return best


def parse_answer_cell(cell_text: str, start_order: int, end_order: int):
    raw = cell_text.replace("\xa0", " ")
    raw = re.sub(r"(?i)\b\d{1,2}\s*[-–]\s*\d{1,2}\s*\.\s*", "", raw)
    raw = re.sub(r"(?i)\b\d{1,2}\s*\.\s*", "", raw)
    parts = []
    for line in raw.splitlines():
        for piece in re.split(r"\s*/\s*", line):
            piece = norm_text(piece)
            if piece:
                parts.append(piece)

    answers = {}
    order = start_order
    for piece in parts:
        range_match = re.match(r"^(\d{1,2})\s*[-–]\s*(\d{1,2})\.?\s*(.+)$", piece)
        if range_match:
            a, b, rest = int(range_match.group(1)), int(range_match.group(2)), range_match.group(3)
            values = [norm_text(x) for x in re.split(r"\s*,\s*", rest) if norm_text(x)]
            for idx, q_order in enumerate(range(a, b + 1)):
                if idx < len(values):
                    answers[q_order] = values[idx]
            order = max(order, b + 1)
            continue

        values = [norm_text(x) for x in re.split(r"\s*,\s*", piece) if norm_text(x)]
        if len(values) > 1 and order + len(values) - 1 <= end_order and all(re.fullmatch(r"[A-Z]+|TRUE|FALSE|NOT GIVEN|YES|NO", v, re.I) for v in values):
            for value in values:
                answers[order] = value
                order += 1
        else:
            answers[order] = piece
            order += 1
    return answers


def extract_explanations(key_path: Path):
    doc = Document(key_path)
    lines = [norm_text(p.text) for p in doc.paragraphs if norm_text(p.text)]
    chunks = {}
    current = None
    buf = []
    for line in lines:
        m = re.match(r"^(\d{1,2})(?:\s*[-–]\s*(\d{1,2}))?\.\s*(.+)?$", line)
        if m:
            if current is not None:
                chunks[current] = "\n".join(buf).strip()
            a = int(m.group(1))
            b = int(m.group(2)) if m.group(2) else a
            current = tuple(range(a, b + 1))
            buf = [line]
        elif current is not None:
            buf.append(line)
    if current is not None:
        chunks[current] = "\n".join(buf).strip()

    explanations = {}
    for orders, text in chunks.items():
        for order in orders:
            explanations[order] = text_to_html(text)
    return explanations


def guess_qset_type(orders, question_texts):
    text = " ".join(question_texts.get(order, "") for order in orders)
    if re.search(r"List of Headings|heading", text, re.I):
        return "MATCHING_HEADINGS"
    if re.search(r"TRUE|FALSE|NOT GIVEN|YES|NO", text, re.I):
        return "SHORT_ANSWER"
    if re.search(r"Choose\s+TWO|Which\s+TWO", text, re.I):
        return "SHORT_ANSWER"
    if re.search(r"\bA\s+.+\bB\s+", text):
        return "SHORT_ANSWER"
    return "SHORT_ANSWER"


def make_question(qid, order, answer, question_text, explanation):
    return {
        "id": qid,
        "order": order,
        "type": "SHORT_ANSWER",
        "text": question_text or f"Question {order}",
        "content": "",
        "options": None,
        "correctAnswer": answer,
        "correctAnswers": None,
        "explanationHtml": explanation or f"<p><strong>Đáp án:</strong> {html.escape(answer)}</p>",
        "locateInfo": None,
        "matchingHeadingParagraph": None,
        "mapPosition": None,
        "audioUrl": None,
        "sampleAnswers": None,
    }


def make_passage_json(test_no: int, passage_no: int, chunk, answers, explanations):
    quiz_id = BASE_ID + test_no * 10 + passage_no
    title = get_title(chunk, passage_no)
    body = passage_body(chunk)
    question_texts = extract_question_texts(chunk)
    start_order = 1 + (passage_no - 1) * 13 if passage_no < 3 else 27
    end_order = 13 if passage_no == 1 else 26 if passage_no == 2 else 40
    orders = list(range(start_order, end_order + 1))

    questions = [
        make_question(
            quiz_id * 100 + order,
            order,
            answers.get(order, ""),
            question_texts.get(order, ""),
            explanations.get(order, ""),
        )
        for order in orders
    ]

    qset = {
        "id": quiz_id * 1000,
        "type": guess_qset_type(orders, question_texts),
        "title": f"Questions {start_order}-{end_order}",
        "instructionHtml": "<p>Answer the questions below.</p>",
        "contentHtml": "",
        "options": None,
        "optionTitle": None,
        "allowReuse": False,
        "maxSelections": 0,
        "image": None,
        "sort": 0,
        "questions": questions,
    }

    return {
        "id": quiz_id,
        "skill": "reading",
        "title": f"[READ VOL 9] Test {test_no} - {title}",
        "type": 1,
        "quizType": 3,
        "durationMin": 60,
        "thumbnail": None,
        "thumbnailUrl": "/assets/thumbs/default.jpg",
        "totalSubmitted": 0,
        "voteCount": 0,
        "isPublic": True,
        "quizCode": f"RV9T{test_no}P{passage_no}",
        "dateUpdated": datetime.now(timezone.utc).isoformat(),
        "tags": [
            {"id": 909, "code": "READ_VOL_9", "title": "READ VOL 9"},
            {"id": 900 + passage_no, "code": f"passage_{passage_no}", "title": f"Passage {passage_no}"},
        ],
        "parts": [
            {
                "id": quiz_id,
                "index": passage_no,
                "title": title,
                "passageHtml": "".join(html_p(line) for line in body),
                "transcriptHtml": None,
                "fileId": None,
                "listenFrom": None,
                "listenTo": None,
                "instruction": None,
                "taskInstruction": None,
                "questionSets": [qset],
                "explanations": [],
            }
        ],
        "explanations": [],
    }


def load_json(path, fallback):
    if not path.exists():
        return deepcopy(fallback)
    return json.loads(path.read_text(encoding="utf-8"))


def main():
    if not SOURCE_DIR.exists():
        raise RuntimeError(f"Missing source dir: {SOURCE_DIR}")

    created_ids = []
    full_tests = load_json(FULL_TEST_INDEX, [])
    normalized_index = load_json(NORMALIZED_INDEX, [])
    normalized_index = [item for item in normalized_index if not (isinstance(item.get("id"), int) and BASE_ID < item["id"] < BASE_ID + 200)]
    full_tests = [item for item in full_tests if not (str(item.get("key", "")).startswith("RV9T") and item.get("skill") == "reading")]

    for test_no in range(1, 10):
        doc_path = SOURCE_DIR / f"VOL 9 TEST {test_no}.docx"
        key_path = KEY_DIR / f"VOL 9 TEST {test_no} KEY.docx"
        passages = split_passages(doc_path)
        if len(passages) != 3:
            raise RuntimeError(f"{doc_path} parsed {len(passages)} passages")

        key_cells = extract_key_cells(key_path)
        explanations = extract_explanations(key_path)
        answer_maps = [
            parse_answer_cell(key_cells[0], 1, 13),
            parse_answer_cell(key_cells[1], 14, 26),
            parse_answer_cell(key_cells[2], 27, 40),
        ]
        for (override_test, order), answer in ANSWER_OVERRIDES.items():
            if override_test == test_no:
                target = 0 if order <= 13 else 1 if order <= 26 else 2
                answer_maps[target][order] = answer

        passage_summaries = []
        for passage_no, chunk in enumerate(passages, start=1):
            quiz = make_passage_json(test_no, passage_no, chunk, answer_maps[passage_no - 1], explanations)
            out_path = NORMALIZED_DIR / f"{quiz['id']}.json"
            out_path.write_text(json.dumps(quiz, ensure_ascii=False, separators=(",", ":")), encoding="utf-8")
            created_ids.append(quiz["id"])
            passage_summaries.append({
                "id": quiz["id"],
                "order": passage_no,
                "title": quiz["parts"][0]["title"],
                "questions": 13 if passage_no < 3 else 14,
            })
            normalized_index.append({
                "id": quiz["id"],
                "title": quiz["title"],
                "parts": 1,
                "questions": 13 if passage_no < 3 else 14,
                "hasPassage": True,
            })

        full_tests.append({
            "key": f"RV9T{test_no}",
            "cambridge": 900 + VOL,
            "test": test_no,
            "skill": "reading",
            "title": f"READ VOL 9 - Test {test_no} (Reading)",
            "passages": passage_summaries,
            "totalQuestions": 40,
            "durationMin": 60,
        })

    NORMALIZED_INDEX.write_text(json.dumps(normalized_index, ensure_ascii=False, indent=2), encoding="utf-8")
    FULL_TEST_INDEX.write_text(json.dumps(full_tests, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Imported {len(created_ids)} passages: {created_ids[0]}..{created_ids[-1]}")


if __name__ == "__main__":
    main()
